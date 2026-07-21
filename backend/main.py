import traceback, json
from fastapi import FastAPI, Depends, HTTPException, Request, WebSocket, WebSocketDisconnect, UploadFile, File, Form
from fastapi.responses import HTMLResponse
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates
from sqlalchemy.orm import Session
from backend.database import engine, get_db, Base, SessionLocal
from backend.models import User, OperationLog, Conversation, Message, Contact
from backend.auth import hash_password, create_token, get_current_user, require_admin
from backend.agent import run_agent
from backend import init_db
from backend.models import Document

init_db()
app = FastAPI()
app.mount("/static", StaticFiles(directory="static"), name="static")
templates = Jinja2Templates(directory="templates")

LOGIN_HTML = """<!DOCTYPE html><html lang="zh"><head><meta charset="UTF-8"><title>AI Office</title><link rel="stylesheet" href="/static/style.css"></head><body>
<div class="auth-container" id="loginBox">
<h2>登录</h2>
<input type="text" id="loginUser" placeholder="用户名" minlength="3" required><br><br>
<input type="password" id="loginPass" placeholder="密码" minlength="6" required><br><br>
<button class="btn-primary" onclick="doLogin()">登录</button>
<p style="text-align:center;margin-top:12px;"><a href="#" onclick="showReg()">注册新账号</a></p>
<p id="loginErr" style="color:red;"></p>
</div>
<div class="auth-container" id="regBox" style="display:none;">
<h2>注册</h2>
<input type="text" id="regUser" placeholder="用户名（3-20位）" minlength="3" maxlength="20" required><br><br>
<input type="password" id="regPass" placeholder="密码（至少6位）" minlength="6" required><br><br>
<input type="password" id="regPass2" placeholder="确认密码"><br><br>
<button class="btn-primary" onclick="doReg()">注册</button>
<p style="text-align:center;margin-top:12px;"><a href="#" onclick="showLogin()">返回登录</a></p>
<p id="regErr" style="color:red;"></p>
</div>
<script>
async function doLogin(){const u=document.getElementById('loginUser').value.trim();const p=document.getElementById('loginPass').value;if(u.length<3){document.getElementById('loginErr').textContent='用户名至少3位';return;}if(p.length<6){document.getElementById('loginErr').textContent='密码至少6位';return;}const r=await fetch('/api/login',{method:'POST',headers:{'Content-Type':'application/json'},body:JSON.stringify({username:u,password:p})});const d=await r.json();if(r.ok){localStorage.setItem('token',d.token);localStorage.setItem('role',d.role);localStorage.setItem('userId',d.user_id);window.location.href='/chat';}else document.getElementById('loginErr').textContent=d.detail;}
async function doReg(){const u=document.getElementById('regUser').value.trim();const pw=document.getElementById('regPass').value;if(u.length<3){document.getElementById('regErr').textContent='用户名至少3位';return;}if(!/^[a-zA-Z0-9_\u4e00-\u9fff]+$/.test(u)){document.getElementById('regErr').textContent='用户名只能包含中英文、数字和下划线';return;}if(pw.length<6){document.getElementById('regErr').textContent='密码至少6位';return;}if(pw!==document.getElementById('regPass2').value){document.getElementById('regErr').textContent='两次密码不一致';return;}const r=await fetch('/api/register',{method:'POST',headers:{'Content-Type':'application/json'},body:JSON.stringify({username:u,password:pw})});const d=await r.json();if(r.ok){alert('注册成功，请登录');showLogin();}else document.getElementById('regErr').textContent=d.detail;}
function showReg(){document.getElementById('loginBox').style.display='none';document.getElementById('regBox').style.display='block';}
function showLogin(){document.getElementById('regBox').style.display='none';document.getElementById('loginBox').style.display='block';}
</script></body></html>"""

@app.get("/", response_class=HTMLResponse)
async def home():
    return HTMLResponse(content=LOGIN_HTML)

@app.get("/register", response_class=HTMLResponse)
async def register_page():
    return HTMLResponse(content=LOGIN_HTML)

@app.post("/api/register")
async def api_register(request: Request, db: Session = Depends(get_db)):
    data = await request.json()
    username = data.get("username", "").strip()
    password = data.get("password", "")
    if not username or not password:
        raise HTTPException(400, "Required")
    if db.query(User).filter(User.username == username).first():
        raise HTTPException(400, "Exists")
    is_first = db.query(User).count() == 0
    user = User(username=username, password=hash_password(password), role="admin" if is_first else "employee")
    db.add(user); db.commit()
    return {"message": "OK"}


@app.get("/api/download/{filename:path}")
async def download_export(filename: str):
    import os
    from fastapi.responses import FileResponse
    exports_dir = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "static", "exports")
    filepath = os.path.join(exports_dir, filename)
    if not os.path.exists(filepath):
        raise HTTPException(404, "File not found")
    return FileResponse(filepath, filename=filename, media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

@app.post("/api/login")
async def api_login(request: Request, db: Session = Depends(get_db)):
    data = await request.json()
    user = db.query(User).filter(User.username == data.get("username","").strip(), User.password == hash_password(data.get("password",""))).first()
    if not user: raise HTTPException(401, "Invalid")
    token = create_token(user.id, user.role)
    return {"token": token, "role": user.role, "username": user.username, "user_id": user.id}
@app.get("/chat", response_class=HTMLResponse)
async def chat_page(request: Request):
    return templates.TemplateResponse(request=request, name="chat.html")

@app.get("/admin", response_class=HTMLResponse)
async def admin_page(request: Request):
    return templates.TemplateResponse(request=request, name="admin.html")

@app.get("/api/admin")
async def api_admin(current_user: User = Depends(require_admin), db: Session = Depends(get_db)):
    from sqlalchemy import func
    logs = db.query(OperationLog).order_by(OperationLog.created_at.desc()).limit(50).all()
    stats = db.query(OperationLog.action, func.count(OperationLog.id).label("count")).group_by(OperationLog.action).all()
    return {"stats": [(s[0], s[1]) for s in stats], "logs": [{"time": log.created_at.strftime('%m-%d %H:%M'), "user": log.user.username, "action": log.action, "detail": log.detail[:50] if log.detail else ""} for log in logs]}

@app.get("/api/conversations")
async def list_conversations(current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    convs = db.query(Conversation).filter(Conversation.user_id == current_user.id).order_by(Conversation.created_at.desc()).all()
    return [{"id": c.id, "title": c.title} for c in convs]

@app.get("/api/conversations/{conv_id}")
async def get_conversation(conv_id: int, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    conv = db.query(Conversation).filter(Conversation.id == conv_id, Conversation.user_id == current_user.id).first()
    if not conv:
        return []
    msgs = db.query(Message).filter(Message.conversation_id == conv_id).all()
    return [{"role": m.role, "content": m.content} for m in msgs]

@app.post("/api/conversations")
async def create_conversation(request: Request, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    data = await request.json()
    conv = Conversation(user_id=current_user.id, title=data.get("title","New")[:30])
    db.add(conv); db.commit()
    return {"id": conv.id}

@app.get("/api/contacts")
async def list_contacts(db: Session = Depends(get_db)):
    contacts = db.query(Contact).all()
    return [{"id": c.id, "name": c.name, "email": c.email, "department": c.department} for c in contacts]

@app.post("/api/contacts")
async def add_contact(request: Request, db: Session = Depends(get_db)):
    data = await request.json()
    c = Contact(name=data.get("name",""), email=data.get("email",""), department=data.get("department",""))
    db.add(c); db.commit()
    return {"id": c.id}

@app.post("/api/admin/promote")
async def promote_user(request: Request, current_user: User = Depends(require_admin), db: Session = Depends(get_db)):
    data = await request.json()
    uid = data.get("user_id")
    user = db.query(User).filter(User.id == uid).first()
    if not user: raise HTTPException(400, "User not found")
    user.role = "admin"
    db.commit()
    return {"message": "User " + user.username + " is now admin"}

@app.websocket("/ws/chat")
async def websocket_chat(websocket: WebSocket):
    await websocket.accept()
    token = websocket.query_params.get("token","")
    if not token:
        await websocket.send_text(json.dumps({"type":"error","content":"Not logged in"}))
        await websocket.close(); return
    import jwt as pyjwt
    from backend.config import SECRET_KEY
    try:
        payload = pyjwt.decode(token, SECRET_KEY, algorithms=["HS256"])
        user_id = payload.get("user_id")
    except pyjwt.PyJWTError:
        await websocket.send_text(json.dumps({"type":"error","content":"Invalid token"}))
        await websocket.close(); return

    await websocket.send_text(json.dumps({"type":"reply","content":"AI Office Agent ready"}))

    try:
        while True:
            data = await websocket.receive_text()
            msg = json.loads(data).get("message","")
            if not msg.strip(): continue
            try:
                # 加载对话历史
                history = []
                conv_id = json.loads(data).get("conv_id", None)
                if conv_id:
                    db3 = SessionLocal()
                    prev = db3.query(Message).filter(Message.conversation_id == conv_id).order_by(Message.created_at.desc()).limit(10).all()
                    db3.close()
                    history = [{"role": m.role, "content": m.content} for m in reversed(prev)]

                ai_reply, tool_results = await run_agent(msg, history=history)
                db_log = SessionLocal()
                conv_id = json.loads(data).get("conv_id", None)
                if conv_id:
                    db_log.add(Message(conversation_id=conv_id, role="user", content=msg))
                    db_log.add(Message(conversation_id=conv_id, role="assistant", content=ai_reply))
                                # 工具名映射
                name_map = {
                    "generate_weekly_report": "生成周报",
                    "summarize_meeting": "会议纪要",
                    "process_excel": "数据处理",
                    "send_notification": "发送通知",
                    "search_knowledge": "搜索知识库",
                    "search_contacts": "搜索通讯录",
                }
                for r in tool_results:
                    raw_name = r.split("]")[0].replace("[","") if "]" in r else "chat"
                    # 只保留结果内容，去掉原始前缀
                    detail = r.split("]",1)[1].strip() if "]" in r else r[:200]
                    log = OperationLog(
                        user_id=user_id,
                        action=name_map.get(raw_name, raw_name),
                        detail=detail[:200]
                    )
                    db_log.add(log)
                db_log.commit(); db_log.close()
                await websocket.send_text(json.dumps({"type":"reply","content":ai_reply,"tools":tool_results}))
            except Exception as e:
                await websocket.send_text(json.dumps({"type":"error","content":f"Error: {str(e)}"}))
    except WebSocketDisconnect:
        pass

@app.get("/api/admin/users")
async def list_users(current_user: User = Depends(require_admin), db: Session = Depends(get_db)):
    users = db.query(User).all()
    return [{"id": u.id, "username": u.username, "role": u.role, "created_at": u.created_at.strftime('%m-%d')} for u in users]

@app.delete("/api/admin/users/{uid}")
async def delete_user(uid: int, current_user: User = Depends(require_admin), db: Session = Depends(get_db)):
    if uid == current_user.id: raise HTTPException(400, "Cannot delete yourself")
    db.query(Message).filter(Message.conversation_id.in_(
        db.query(Conversation.id).filter(Conversation.user_id == uid)
    )).delete(synchronize_session=False)
    db.query(Conversation).filter(Conversation.user_id == uid).delete(synchronize_session=False)
    db.query(OperationLog).filter(OperationLog.user_id == uid).delete(synchronize_session=False)
    db.query(User).filter(User.id == uid).delete(synchronize_session=False)
    db.commit()
    return {"message": "Deleted"}

@app.get("/api/conversations/{conv_id}/export")
async def export_conversation(conv_id: int, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    msgs = db.query(Message).filter(Message.conversation_id == conv_id).order_by(Message.created_at).all()
    conv = db.query(Conversation).filter(Conversation.id == conv_id).first()
    if not conv or conv.user_id != current_user.id:
        return {"text": "无权限或对话不存在"}
    text = f"Conversation: {conv.title if conv else 'Unknown'}\n{'='*40}\n\n"
    for m in msgs:
        role = "User" if m.role == "user" else "AI"
        text += f"[{role}] {m.created_at.strftime('%m-%d %H:%M')}\n{m.content}\n\n"
    return {"text": text}

@app.post("/api/documents")
async def add_document(request: Request, current_user: User = Depends(require_admin), db: Session = Depends(get_db)):
    data = await request.json()
    doc = Document(title=data.get("title",""), content=data.get("content",""))
    db.add(doc); db.commit()
    return {"id": doc.id}

@app.get("/api/documents")
async def list_documents(db: Session = Depends(get_db)):
    docs = db.query(Document).all()
    return [{"id": d.id, "title": d.title} for d in docs]

@app.post("/api/documents/upload")
async def upload_document(
    file: UploadFile = File(...),
    current_user: User = Depends(get_current_user)
):
    content = ""
    filename = file.filename or "unknown"

    if filename.endswith(".txt"):
        content = (await file.read()).decode("utf-8", errors="ignore")
    elif filename.endswith(".pdf"):
        import io
        from pypdf import PdfReader
        data = await file.read()
        reader = PdfReader(io.BytesIO(data))
        for page in reader.pages:
            content += page.extract_text() or ""
    elif filename.endswith(".docx"):
        import io
        from docx import Document as DocxDoc
        data = await file.read()
        doc = DocxDoc(io.BytesIO(data))
        content = "\n".join([p.text for p in doc.paragraphs])
    elif filename.endswith(".xlsx"):
        import io
        from openpyxl import load_workbook
        data = await file.read()
        wb = load_workbook(io.BytesIO(data), data_only=True)
        lines = []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            lines.append(f"=== Sheet: {sheet_name} ===")
            for row in ws.iter_rows(values_only=True):
                line = " | ".join([str(c) if c is not None else "" for c in row])
                lines.append(line)
        content = "\n".join(lines)
    else:
        raise HTTPException(400, "Unsupported format. Use .txt, .pdf, .docx, or .xlsx")

    if not content.strip():
        raise HTTPException(400, "No text content found in file")

    db2 = SessionLocal()
    doc = Document(title=filename, content=content[:10000])
    db2.add(doc); db2.commit()
    doc_id = doc.id
    db2.close()
    return {"id": doc_id, "title": filename}

from apscheduler.schedulers.asyncio import AsyncIOScheduler

scheduler = AsyncIOScheduler()

async def daily_reminder():
    import os, smtplib, asyncio
    from email.mime.text import MIMEText
    try:
        smtp_host = os.getenv("SMTP_HOST", "smtp.qq.com")
        smtp_user = os.getenv("SMTP_USER", "")
        smtp_pass = os.getenv("SMTP_PASS", "")
        if not smtp_user: return False, "SMTP_USER not set"
        if not smtp_pass: return False, "SMTP_PASS not set"

        msg = MIMEText("Reminder: submit weekly report.", "plain", "utf-8")
        msg["From"] = smtp_user
        msg["To"] = smtp_user
        msg["Subject"] = "Weekly Report Reminder"

        def _send():
            s = smtplib.SMTP_SSL(smtp_host, 465, timeout=10)
            s.login(smtp_user, smtp_pass)
            s.send_message(msg)
            s.quit()
        await asyncio.get_event_loop().run_in_executor(None, _send)
        return True, "Sent to " + smtp_user
    except Exception as e:
        return False, str(e)

@app.post("/api/reminder/test")
async def test_reminder():
    ok, msg = await daily_reminder()
    return {"ok": ok, "message": msg}

@app.delete("/api/documents/{doc_id}")
async def delete_document(doc_id: int, current_user: User = Depends(require_admin), db: Session = Depends(get_db)):
    doc = db.query(Document).filter(Document.id == doc_id).first()
    if not doc: raise HTTPException(404, "Not found")
    db.delete(doc); db.commit()
    return {"message": "Deleted"}

@app.put("/api/documents/{doc_id}")
async def update_document(doc_id: int, request: Request, current_user: User = Depends(require_admin), db: Session = Depends(get_db)):
    data = await request.json()
    doc = db.query(Document).filter(Document.id == doc_id).first()
    if not doc: raise HTTPException(404, "Not found")
    doc.title = data.get("title", doc.title)
    doc.content = data.get("content", doc.content)
    db.commit()
    return {"message": "Updated"}    

if __name__ == "__main__":
    import uvicorn
    scheduler.add_job(daily_reminder, "cron", hour=23, minute=18)
    scheduler.start()
    uvicorn.run("backend.main:app", host="127.0.0.1", port=5000, reload=False)
